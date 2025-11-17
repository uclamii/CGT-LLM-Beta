#!/usr/bin/env python3
"""
RAG Chatbot Implementation for CGT-LLM-Beta with Vector Database
Production-ready local RAG system with ChromaDB and MPS acceleration for Apple Silicon
"""

import argparse
import csv
import json
import logging
import os
import re
import sys
import time
import hashlib
from pathlib import Path
from typing import List, Tuple, Dict, Any, Optional, Union
from dataclasses import dataclass
from collections import defaultdict

import textstat

import torch
import numpy as np
import pandas as pd
from tqdm import tqdm

# Optional imports with graceful fallbacks
try:
    import chromadb
    from chromadb.config import Settings
    CHROMADB_AVAILABLE = True
except ImportError:
    CHROMADB_AVAILABLE = False
    print("Warning: chromadb not available. Install with: pip install chromadb")

try:
    from sentence_transformers import SentenceTransformer
    SENTENCE_TRANSFORMERS_AVAILABLE = True
except ImportError:
    SENTENCE_TRANSFORMERS_AVAILABLE = False
    print("Warning: sentence-transformers not available. Install with: pip install sentence-transformers")

try:
    import pypdf
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("Warning: pypdf not available. PDF files will be skipped.")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. DOCX files will be skipped.")

try:
    from rank_bm25 import BM25Okapi
    BM25_AVAILABLE = True
except ImportError:
    BM25_AVAILABLE = False
    print("Warning: rank-bm25 not available. BM25 retrieval disabled.")

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('rag_bot.log')
    ]
)
logger = logging.getLogger(__name__)


@dataclass
class Document:
    """Represents a document with metadata"""
    filename: str
    content: str
    filepath: str
    file_type: str
    chunk_count: int = 0
    file_hash: str = ""


@dataclass
class Chunk:
    """Represents a text chunk with metadata"""
    text: str
    filename: str
    chunk_id: int
    total_chunks: int
    start_pos: int
    end_pos: int
    metadata: Dict[str, Any]
    chunk_hash: str = ""


class VectorRetriever:
    """ChromaDB-based vector retrieval"""
    
    def __init__(self, collection_name: str = "cgt_documents", persist_directory: str = "./chroma_db"):
        if not CHROMADB_AVAILABLE:
            raise ImportError("ChromaDB is required for vector retrieval")
        
        self.collection_name = collection_name
        self.persist_directory = persist_directory
        
        # Initialize ChromaDB client
        self.client = chromadb.PersistentClient(path=persist_directory)
        
        # Get or create collection
        try:
            self.collection = self.client.get_collection(name=collection_name)
            logger.info(f"Loaded existing collection '{collection_name}' with {self.collection.count()} documents")
        except:
            self.collection = self.client.create_collection(
                name=collection_name,
                metadata={"description": "CGT-LLM-Beta document collection"}
            )
            logger.info(f"Created new collection '{collection_name}'")
        
        # Initialize embedding model
        if SENTENCE_TRANSFORMERS_AVAILABLE:
            self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
            logger.info("Loaded sentence-transformers embedding model")
        else:
            self.embedding_model = None
            logger.warning("Sentence-transformers not available, using ChromaDB default embeddings")
    
    def add_documents(self, chunks: List[Chunk]) -> None:
        """Add document chunks to the vector database"""
        if not chunks:
            return
        
        logger.info(f"Adding {len(chunks)} chunks to vector database...")
        
        # Prepare data for ChromaDB
        documents = []
        metadatas = []
        ids = []
        
        for chunk in chunks:
            chunk_id = f"{chunk.filename}_{chunk.chunk_id}"
            documents.append(chunk.text)
            
            metadata = {
                "filename": chunk.filename,
                "chunk_id": chunk.chunk_id,
                "total_chunks": chunk.total_chunks,
                "start_pos": chunk.start_pos,
                "end_pos": chunk.end_pos,
                "chunk_hash": chunk.chunk_hash,
                **chunk.metadata
            }
            metadatas.append(metadata)
            ids.append(chunk_id)
        
        # Add to collection
        try:
            self.collection.add(
                documents=documents,
                metadatas=metadatas,
                ids=ids
            )
            logger.info(f"Successfully added {len(chunks)} chunks to vector database")
        except Exception as e:
            logger.error(f"Error adding documents to vector database: {e}")
    
    def search(self, query: str, k: int = 5) -> List[Tuple[Chunk, float]]:
        """Search for similar chunks using vector similarity"""
        try:
            # Perform vector search
            results = self.collection.query(
                query_texts=[query],
                n_results=k
            )
            
            chunks_with_scores = []
            if results['documents'] and results['documents'][0]:
                for i, (doc, metadata, distance) in enumerate(zip(
                    results['documents'][0],
                    results['metadatas'][0],
                    results['distances'][0]
                )):
                    # Convert distance to similarity score (ChromaDB uses cosine distance)
                    similarity_score = 1 - distance
                    
                    chunk = Chunk(
                        text=doc,
                        filename=metadata['filename'],
                        chunk_id=metadata['chunk_id'],
                        total_chunks=metadata['total_chunks'],
                        start_pos=metadata['start_pos'],
                        end_pos=metadata['end_pos'],
                        metadata={k: v for k, v in metadata.items() 
                                if k not in ['filename', 'chunk_id', 'total_chunks', 'start_pos', 'end_pos', 'chunk_hash']},
                        chunk_hash=metadata.get('chunk_hash', '')
                    )
                    chunks_with_scores.append((chunk, similarity_score))
            
            return chunks_with_scores
            
        except Exception as e:
            logger.error(f"Error searching vector database: {e}")
            return []
    
    def get_collection_stats(self) -> Dict[str, Any]:
        """Get statistics about the collection"""
        try:
            count = self.collection.count()
            return {
                "total_chunks": count,
                "collection_name": self.collection_name,
                "persist_directory": self.persist_directory
            }
        except Exception as e:
            logger.error(f"Error getting collection stats: {e}")
            return {}


class RAGBot:
    """Main RAG chatbot class with vector database"""
    
    def __init__(self, args):
        self.args = args
        self.device = self._setup_device()
        self.model = None
        self.tokenizer = None
        self.vector_retriever = None
        
        # Load model
        self._load_model()
        
        # Initialize vector retriever
        self._setup_vector_retriever()
    
    def _setup_device(self) -> str:
        """Setup device with MPS support for Apple Silicon"""
        if torch.backends.mps.is_available():
            device = "mps"
            logger.info("Using device: mps (Apple Silicon)")
        elif torch.cuda.is_available():
            device = "cuda"
            logger.info("Using device: cuda")
        else:
            device = "cpu"
            logger.info("Using device: cpu")
        
        return device
    
    def _load_model(self):
        """Load the specified LLM model and tokenizer"""
        try:
            model_name = self.args.model
            logger.info(f"Loading model: {model_name}...")
            from transformers import AutoTokenizer, AutoModelForCausalLM
            
            # Load tokenizer
            self.tokenizer = AutoTokenizer.from_pretrained(
                model_name,
                trust_remote_code=True
            )
            
            # Determine appropriate torch dtype based on device and model
            # Use float16 for MPS/CUDA, float32 for CPU
            # Some models work better with bfloat16
            if self.device == "mps":
                torch_dtype = torch.float16
            elif self.device == "cuda":
                torch_dtype = torch.float16
            else:
                torch_dtype = torch.float32
            
            # Load model with appropriate settings
            model_kwargs = {
                "torch_dtype": torch_dtype,
                "trust_remote_code": True,
            }
            
            # For MPS, use device_map; for CUDA, let it auto-detect
            if self.device == "mps":
                model_kwargs["device_map"] = self.device
            elif self.device == "cuda":
                model_kwargs["device_map"] = "auto"
            # For CPU, don't specify device_map
            
            self.model = AutoModelForCausalLM.from_pretrained(
                model_name,
                **model_kwargs
            )
            
            # Move to device if not using device_map
            if self.device == "cpu":
                self.model = self.model.to(self.device)
            
            # Set pad token if not already set
            if self.tokenizer.pad_token is None:
                if self.tokenizer.eos_token is not None:
                    self.tokenizer.pad_token = self.tokenizer.eos_token
                else:
                    # Some models might need a different approach
                    self.tokenizer.add_special_tokens({'pad_token': '[PAD]'})
            
            logger.info(f"Model {model_name} loaded successfully on {self.device}")
            
        except Exception as e:
            logger.error(f"Failed to load model {self.args.model}: {e}")
            logger.error("Make sure the model name is correct and you have access to it on HuggingFace")
            logger.error("For private models, ensure you're logged in: huggingface-cli login")
            sys.exit(2)
    
    def _setup_vector_retriever(self):
        """Setup the vector retriever"""
        try:
            self.vector_retriever = VectorRetriever(
                collection_name="cgt_documents",
                persist_directory=self.args.vector_db_dir
            )
            logger.info("Vector retriever initialized successfully")
        except Exception as e:
            logger.error(f"Failed to setup vector retriever: {e}")
            sys.exit(2)
    
    def _calculate_file_hash(self, filepath: str) -> str:
        """Calculate hash of file for change detection"""
        try:
            with open(filepath, 'rb') as f:
                return hashlib.md5(f.read()).hexdigest()
        except:
            return ""
    
    def _calculate_chunk_hash(self, text: str) -> str:
        """Calculate hash of chunk text"""
        return hashlib.md5(text.encode('utf-8')).hexdigest()
    
    def load_corpus(self, data_dir: str) -> List[Document]:
        """Load all documents from the data directory"""
        logger.info(f"Loading corpus from {data_dir}")
        documents = []
        data_path = Path(data_dir)
        
        if not data_path.exists():
            logger.error(f"Data directory {data_dir} does not exist")
            sys.exit(1)
        
        # Supported file extensions
        supported_extensions = {'.txt', '.md', '.json', '.csv'}
        if PDF_AVAILABLE:
            supported_extensions.add('.pdf')
        if DOCX_AVAILABLE:
            supported_extensions.add('.docx')
            supported_extensions.add('.doc')
        
        # Find all files recursively
        files = []
        for ext in supported_extensions:
            files.extend(data_path.rglob(f"*{ext}"))
        
        logger.info(f"Found {len(files)} files to process")
        
        # Process files with progress bar
        for file_path in tqdm(files, desc="Loading documents"):
            try:
                content = self._read_file(file_path)
                if content.strip():  # Only add non-empty documents
                    file_hash = self._calculate_file_hash(file_path)
                    doc = Document(
                        filename=file_path.name,
                        content=content,
                        filepath=str(file_path),
                        file_type=file_path.suffix.lower(),
                        file_hash=file_hash
                    )
                    documents.append(doc)
                    logger.debug(f"Loaded {file_path.name} ({len(content)} chars)")
                else:
                    logger.warning(f"Skipping empty file: {file_path.name}")
                    
            except Exception as e:
                logger.error(f"Failed to load {file_path.name}: {e}")
                continue
        
        logger.info(f"Successfully loaded {len(documents)} documents")
        return documents
    
    def _read_file(self, file_path: Path) -> str:
        """Read content from various file types"""
        suffix = file_path.suffix.lower()
        
        try:
            if suffix == '.txt':
                return file_path.read_text(encoding='utf-8')
            
            elif suffix == '.md':
                return file_path.read_text(encoding='utf-8')
            
            elif suffix == '.json':
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    if isinstance(data, dict):
                        return json.dumps(data, indent=2)
                    else:
                        return str(data)
            
            elif suffix == '.csv':
                df = pd.read_csv(file_path)
                return df.to_string()
            
            elif suffix == '.pdf' and PDF_AVAILABLE:
                text = ""
                with open(file_path, 'rb') as f:
                    pdf_reader = pypdf.PdfReader(f)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                return text
            
            elif suffix in ['.docx', '.doc'] and DOCX_AVAILABLE:
                doc = Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text
            
            else:
                logger.warning(f"Unsupported file type: {suffix}")
                return ""
                
        except Exception as e:
            logger.error(f"Error reading {file_path}: {e}")
            return ""
    
    def chunk_documents(self, docs: List[Document], chunk_size: int, overlap: int) -> List[Chunk]:
        """Chunk documents into smaller pieces"""
        logger.info(f"Chunking {len(docs)} documents (size={chunk_size}, overlap={overlap})")
        chunks = []
        
        for doc in docs:
            doc_chunks = self._chunk_text(
                doc.content, 
                doc.filename, 
                chunk_size, 
                overlap
            )
            chunks.extend(doc_chunks)
            
            # Update document metadata
            doc.chunk_count = len(doc_chunks)
        
        logger.info(f"Created {len(chunks)} chunks from {len(docs)} documents")
        return chunks
    
    def _chunk_text(self, text: str, filename: str, chunk_size: int, overlap: int) -> List[Chunk]:
        """Split text into overlapping chunks"""
        # Clean text
        text = re.sub(r'\s+', ' ', text.strip())
        
        # Simple token-based chunking (approximate)
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk_words = words[i:i + chunk_size]
            chunk_text = ' '.join(chunk_words)
            
            if chunk_text.strip():
                chunk_hash = self._calculate_chunk_hash(chunk_text)
                chunk = Chunk(
                    text=chunk_text,
                    filename=filename,
                    chunk_id=len(chunks),
                    total_chunks=0,  # Will be updated later
                    start_pos=i,
                    end_pos=i + len(chunk_words),
                    metadata={
                        'word_count': len(chunk_words),
                        'char_count': len(chunk_text)
                    },
                    chunk_hash=chunk_hash
                )
                chunks.append(chunk)
        
        # Update total_chunks for each chunk
        for chunk in chunks:
            chunk.total_chunks = len(chunks)
        
        return chunks
    
    def build_or_update_index(self, chunks: List[Chunk], force_rebuild: bool = False) -> None:
        """Build or update the vector index"""
        if not chunks:
            logger.warning("No chunks provided for indexing")
            return
        
        # Check if we need to rebuild
        collection_stats = self.vector_retriever.get_collection_stats()
        existing_count = collection_stats.get('total_chunks', 0)
        
        if existing_count > 0 and not force_rebuild:
            logger.info(f"Vector database already contains {existing_count} chunks. Use --force-rebuild to rebuild.")
            return
        
        if force_rebuild and existing_count > 0:
            logger.info("Force rebuild requested. Clearing existing collection...")
            try:
                self.client.delete_collection(self.vector_retriever.collection_name)
                self.vector_retriever.collection = self.client.create_collection(
                    name=self.vector_retriever.collection_name,
                    metadata={"description": "CGT-LLM-Beta document collection"}
                )
            except Exception as e:
                logger.error(f"Error clearing collection: {e}")
        
        # Add chunks to vector database
        self.vector_retriever.add_documents(chunks)
        
        logger.info("Vector index built successfully")
    
    def retrieve(self, query: str, k: int) -> List[Chunk]:
        """Retrieve relevant chunks for a query using vector search"""
        results = self.vector_retriever.search(query, k)
        chunks = [chunk for chunk, score in results]
        
        if self.args.verbose:
            logger.info(f"Retrieved {len(chunks)} chunks for query: {query[:50]}...")
            for i, (chunk, score) in enumerate(results):
                logger.info(f"  {i+1}. {chunk.filename} (score: {score:.3f})")
        
        return chunks
    
    def retrieve_with_scores(self, query: str, k: int) -> Tuple[List[Chunk], List[float]]:
        """Retrieve relevant chunks with similarity scores
        
        Returns:
            Tuple of (chunks, scores) where scores are similarity scores for each chunk
        """
        results = self.vector_retriever.search(query, k)
        chunks = [chunk for chunk, score in results]
        scores = [score for chunk, score in results]
        
        if self.args.verbose:
            logger.info(f"Retrieved {len(chunks)} chunks for query: {query[:50]}...")
            for i, (chunk, score) in enumerate(results):
                logger.info(f"  {i+1}. {chunk.filename} (score: {score:.3f})")
        
        return chunks, scores
    
    def format_prompt(self, context_chunks: List[Chunk], question: str) -> str:
        """Format the prompt with context and question, ensuring it fits within token limits"""
        context_parts = []
        for chunk in context_chunks:
            context_parts.append(f"{chunk.text}")
        
        context = "\n".join(context_parts)
        
        # Try to use the tokenizer's chat template if available
        if hasattr(self.tokenizer, 'apply_chat_template') and self.tokenizer.chat_template is not None:
            try:
                messages = [
                    {"role": "system", "content": "You are a helpful medical assistant. Answer questions based on the provided context. Be specific and informative."},
                    {"role": "user", "content": f"Context: {context}\n\nQuestion: {question}"}
                ]
                base_prompt = self.tokenizer.apply_chat_template(
                    messages,
                    tokenize=False,
                    add_generation_prompt=True
                )
            except Exception as e:
                logger.warning(f"Failed to use chat template, falling back to manual format: {e}")
                base_prompt = self._format_prompt_manual(context, question)
        else:
            # Fall back to manual formatting (for Llama models)
            base_prompt = self._format_prompt_manual(context, question)
        
        # Check if prompt is too long and truncate context if needed
        max_context_tokens = 1200  # Leave room for generation
        try:
            tokenized = self.tokenizer(base_prompt, return_tensors="pt")
            current_tokens = tokenized['input_ids'].shape[1]
        except Exception as e:
            logger.warning(f"Tokenization error, using base prompt as-is: {e}")
            return base_prompt
        
        if current_tokens > max_context_tokens:
            # Truncate context to fit within limits
            try:
                context_tokens = self.tokenizer(context, return_tensors="pt")['input_ids'].shape[1]
                available_tokens = max_context_tokens - (current_tokens - context_tokens)
                
                if available_tokens > 0:
                    # Truncate context to fit
                    truncated_context = self.tokenizer.decode(
                        self.tokenizer(context, return_tensors="pt", truncation=True, max_length=available_tokens)['input_ids'][0],
                        skip_special_tokens=True
                    )
                    
                    # Reformat with truncated context
                    if hasattr(self.tokenizer, 'apply_chat_template') and self.tokenizer.chat_template is not None:
                        try:
                            messages = [
                                {"role": "system", "content": "You are a helpful medical assistant. Answer questions based on the provided context. Be specific and informative."},
                                {"role": "user", "content": f"Context: {truncated_context}\n\nQuestion: {question}"}
                            ]
                            prompt = self.tokenizer.apply_chat_template(
                                messages,
                                tokenize=False,
                                add_generation_prompt=True
                            )
                        except:
                            prompt = self._format_prompt_manual(truncated_context, question)
                    else:
                        prompt = self._format_prompt_manual(truncated_context, question)
                else:
                    # If even basic prompt is too long, use minimal format
                    prompt = self._format_prompt_manual(context[:500] + "...", question)
            except Exception as e:
                logger.warning(f"Error truncating context: {e}, using base prompt")
                prompt = base_prompt
        else:
            prompt = base_prompt
            
        return prompt
    
    def _format_prompt_manual(self, context: str, question: str) -> str:
        """Manual prompt formatting for models without chat templates (e.g., Llama)"""
        return f"""<|begin_of_text|><|start_header_id|>system<|end_header_id|>

You are a helpful medical assistant. Answer questions based on the provided context. Be specific and informative.<|eot_id|><|start_header_id|>user<|end_header_id|>

Context: {context}

Question: {question}<|eot_id|><|start_header_id|>assistant<|end_header_id|>

"""
    
    def format_improved_prompt(self, context_chunks: List[Chunk], question: str) -> Tuple[str, str]:
        """Format an improved prompt with better tone, structure, and medical appropriateness
        
        Returns:
            Tuple of (prompt, prompt_text) where prompt_text is the system prompt instructions
        """
        context_parts = []
        for chunk in context_chunks:
            context_parts.append(f"{chunk.text}")
        
        context = "\n".join(context_parts)
        
        # Improved prompt with all the feedback incorporated
        improved_prompt_text = """Provide a concise, neutral, and informative answer based on the provided medical context. 

CRITICAL GUIDELINES:
- Format your response as clear, well-structured sentences and paragraphs
- Be concise and direct - focus on answering the specific question asked
- Use neutral, factual language - do NOT tell the questioner how to feel (avoid phrases like 'don't worry', 'the good news is', etc.)
- Do NOT use leading or coercive language - present information neutrally to preserve patient autonomy
- Do NOT make specific medical recommendations - instead state that management decisions should be made with a healthcare provider
- Use third-person voice only - never claim to be a medical professional or assistant
- Use consistent terminology: use 'children' (not 'offspring') consistently
- Do NOT include hypothetical examples with specific names (e.g., avoid 'Aunt Jenna' or similar)
- Include important distinctions when relevant (e.g., somatic vs. germline variants, reproductive risks)
- When citing sources, be consistent - always specify which guidelines or sources when mentioned
- Remove any formatting markers like asterisks (*) or bold markers
- Do NOT include phrases like 'Here's a rewritten version' - just provide the answer directly

If the question asks about medical management, screening, or interventions, conclude with: 'Management recommendations are individualized and should be discussed with a healthcare provider or genetic counselor.'"""
        
        # Try to use the tokenizer's chat template if available
        if hasattr(self.tokenizer, 'apply_chat_template') and self.tokenizer.chat_template is not None:
            try:
                messages = [
                    {"role": "system", "content": improved_prompt_text},
                    {"role": "user", "content": f"Context: {context}\n\nQuestion: {question}"}
                ]
                base_prompt = self.tokenizer.apply_chat_template(
                    messages,
                    tokenize=False,
                    add_generation_prompt=True
                )
            except Exception as e:
                logger.warning(f"Failed to use chat template for improved prompt, falling back to manual format: {e}")
                base_prompt = self._format_improved_prompt_manual(context, question, improved_prompt_text)
        else:
            # Fall back to manual formatting (for Llama models)
            base_prompt = self._format_improved_prompt_manual(context, question, improved_prompt_text)
        
        # Check if prompt is too long and truncate context if needed
        max_context_tokens = 1200  # Leave room for generation
        try:
            tokenized = self.tokenizer(base_prompt, return_tensors="pt")
            current_tokens = tokenized['input_ids'].shape[1]
        except Exception as e:
            logger.warning(f"Tokenization error for improved prompt, using base prompt as-is: {e}")
            return base_prompt, improved_prompt_text
        
        if current_tokens > max_context_tokens:
            # Truncate context to fit within limits
            try:
                context_tokens = self.tokenizer(context, return_tensors="pt")['input_ids'].shape[1]
                available_tokens = max_context_tokens - (current_tokens - context_tokens)
                
                if available_tokens > 0:
                    # Truncate context to fit
                    truncated_context = self.tokenizer.decode(
                        self.tokenizer(context, return_tensors="pt", truncation=True, max_length=available_tokens)['input_ids'][0],
                        skip_special_tokens=True
                    )
                    
                    # Reformat with truncated context
                    if hasattr(self.tokenizer, 'apply_chat_template') and self.tokenizer.chat_template is not None:
                        try:
                            messages = [
                                {"role": "system", "content": improved_prompt_text},
                                {"role": "user", "content": f"Context: {truncated_context}\n\nQuestion: {question}"}
                            ]
                            prompt = self.tokenizer.apply_chat_template(
                                messages,
                                tokenize=False,
                                add_generation_prompt=True
                            )
                        except:
                            prompt = self._format_improved_prompt_manual(truncated_context, question, improved_prompt_text)
                    else:
                        prompt = self._format_improved_prompt_manual(truncated_context, question, improved_prompt_text)
                else:
                    # If even basic prompt is too long, use minimal format
                    prompt = self._format_improved_prompt_manual(context[:500] + "...", question, improved_prompt_text)
            except Exception as e:
                logger.warning(f"Error truncating context for improved prompt: {e}, using base prompt")
                prompt = base_prompt
        else:
            prompt = base_prompt
            
        return prompt, improved_prompt_text
    
    def _format_improved_prompt_manual(self, context: str, question: str, improved_prompt_text: str) -> str:
        """Manual prompt formatting for improved prompts (for models without chat templates)"""
        return f"""<|begin_of_text|><|start_header_id|>system<|end_header_id|>

{improved_prompt_text}<|eot_id|><|start_header_id|>user<|end_header_id|>

Context: {context}

Question: {question}<|eot_id|><|start_header_id|>assistant<|end_header_id|>

"""
    
    def generate_answer(self, prompt: str, **gen_kwargs) -> str:
        """Generate answer using the language model"""
        try:
            if self.args.verbose:
                logger.info(f"Full prompt (first 500 chars): {prompt[:500]}...")
            
            # Tokenize input with more conservative limit to leave room for generation
            inputs = self.tokenizer(prompt, return_tensors="pt", truncation=True, max_length=1500)
            inputs = {k: v.to(self.device) for k, v in inputs.items()}
            
            if self.args.verbose:
                logger.info(f"Input tokens: {inputs['input_ids'].shape}")
            
            # Generate
            with torch.no_grad():
                outputs = self.model.generate(
                    **inputs,
                    max_new_tokens=gen_kwargs.get('max_new_tokens', 512),
                    temperature=gen_kwargs.get('temperature', 0.7),
                    top_p=gen_kwargs.get('top_p', 0.95),
                    repetition_penalty=gen_kwargs.get('repetition_penalty', 1.05),
                    do_sample=True,
                    pad_token_id=self.tokenizer.eos_token_id,
                    eos_token_id=self.tokenizer.eos_token_id,
                    use_cache=True,
                    num_beams=1
                )
            
            # Decode response without skipping special tokens to preserve full length
            response = self.tokenizer.decode(outputs[0], skip_special_tokens=False)
            
            if self.args.verbose:
                logger.info(f"Full response (first 1000 chars): {response[:1000]}...")
                logger.info(f"Looking for 'Answer:' in response: {'Answer:' in response}")
                if "Answer:" in response:
                    answer_part = response.split("Answer:")[-1]
                    logger.info(f"Answer part (first 200 chars): {answer_part[:200]}...")
                
                # Debug: Show the full response to understand the structure
                logger.info(f"Full response length: {len(response)}")
                logger.info(f"Prompt length: {len(prompt)}")
                logger.info(f"Response after prompt (first 500 chars): {response[len(prompt):][:500]}...")
            
            # Extract the answer more robustly by looking for the end of the prompt
            # Find the actual end of the prompt in the response
            prompt_end_marker = "<|start_header_id|>assistant<|end_header_id|>\n\n"
            if prompt_end_marker in response:
                answer = response.split(prompt_end_marker)[-1].strip()
            else:
                # Fallback to character-based extraction
                answer = response[len(prompt):].strip()
            
            if self.args.verbose:
                logger.info(f"Full LLM output (first 200 chars): {answer[:200]}...")
                logger.info(f"Full LLM output length: {len(answer)} characters")
                logger.info(f"Full LLM output (last 200 chars): ...{answer[-200:]}")
            
            # Only do minimal cleanup to preserve the full response
            # Remove special tokens that might interfere with display, but preserve content
            if "<|start_header_id|>" in answer:
                # Only remove if it's at the very end
                if answer.endswith("<|start_header_id|>"):
                    answer = answer[:-len("<|start_header_id|>")].strip()
            if "<|eot_id|>" in answer:
                # Only remove if it's at the very end
                if answer.endswith("<|eot_id|>"):
                    answer = answer[:-len("<|eot_id|>")].strip()
            if "<|end_of_text|>" in answer:
                # Only remove if it's at the very end
                if answer.endswith("<|end_of_text|>"):
                    answer = answer[:-len("<|end_of_text|>")].strip()
            
            # Final validation - only reject if completely empty
            if not answer or len(answer) < 3:
                answer = "I don't know."
            
            if self.args.verbose:
                logger.info(f"Final answer: '{answer}'")
            
            return answer
            
        except Exception as e:
            logger.error(f"Generation error: {e}")
            return "I encountered an error while generating the answer."
    
    def process_questions(self, questions_path: str, **kwargs) -> List[Tuple[str, str, str, str, float, str, float, str, float, str, str]]:
        """Process all questions and generate answers with multiple readability levels
        
        Returns:
            List of tuples: (question, answer, sources, question_group, original_flesch, 
                            middle_school_answer, middle_school_flesch, 
                            high_school_answer, high_school_flesch, improved_answer, similarity_scores)
        """
        logger.info(f"Processing questions from {questions_path}")
        
        # Load questions
        try:
            with open(questions_path, 'r', encoding='utf-8') as f:
                questions = [line.strip() for line in f if line.strip()]
        except Exception as e:
            logger.error(f"Failed to load questions: {e}")
            sys.exit(1)
        
        logger.info(f"Found {len(questions)} questions to process")
        
        qa_pairs = []
        
        # Get the improved prompt text for CSV header by calling format_improved_prompt with empty chunks
        # This will give us the prompt text without actually generating
        _, improved_prompt_text = self.format_improved_prompt([], "")
        
        # Initialize CSV file with headers
        self.write_csv([], kwargs.get('output_file', 'results.csv'), append=False, improved_prompt_text=improved_prompt_text)
        
        # Process each question
        for i, question in enumerate(tqdm(questions, desc="Processing questions")):
            logger.info(f"Question {i+1}/{len(questions)}: {question[:50]}...")
            
            try:
                # Categorize question
                question_group = self._categorize_question(question)
                
                # Retrieve relevant chunks with similarity scores
                context_chunks, similarity_scores = self.retrieve_with_scores(question, self.args.k)
                
                # Format similarity scores as a string (comma-separated, 3 decimal places)
                similarity_scores_str = ", ".join([f"{score:.3f}" for score in similarity_scores]) if similarity_scores else "0.000"
                
                if not context_chunks:
                    answer = "I don't know."
                    sources = "No sources found"
                    middle_school_answer = "I don't know."
                    high_school_answer = "I don't know."
                    improved_answer = "I don't know."
                    original_flesch = 0.0
                    middle_school_flesch = 0.0
                    high_school_flesch = 0.0
                    similarity_scores_str = "0.000"
                else:
                    # Format original prompt
                    prompt = self.format_prompt(context_chunks, question)
                    
                    # Generate original answer
                    start_time = time.time()
                    answer = self.generate_answer(prompt, **kwargs)
                    gen_time = time.time() - start_time
                    
                    # Generate improved answer
                    improved_prompt, _ = self.format_improved_prompt(context_chunks, question)
                    improved_start = time.time()
                    improved_answer = self.generate_answer(improved_prompt, **kwargs)
                    improved_time = time.time() - improved_start
                    
                    # Clean up improved answer - remove unwanted phrases and formatting
                    improved_answer = self._clean_improved_answer(improved_answer)
                    logger.info(f"Improved answer generated in {improved_time:.2f}s")
                    
                    # Extract source documents
                    sources = self._extract_sources(context_chunks)
                    
                    # Calculate original answer Flesch score
                    try:
                        original_flesch = textstat.flesch_kincaid_grade(answer)
                    except:
                        original_flesch = 0.0
                    
                    # Generate middle school version
                    readability_start = time.time()
                    middle_school_answer, middle_school_flesch = self.enhance_readability(answer, "middle_school")
                    readability_time = time.time() - readability_start
                    logger.info(f"Middle school readability in {readability_time:.2f}s")
                    
                    # Generate high school version
                    readability_start = time.time()
                    high_school_answer, high_school_flesch = self.enhance_readability(answer, "high_school")
                    readability_time = time.time() - readability_start
                    logger.info(f"High school readability in {readability_time:.2f}s")
                    
                    logger.info(f"Generated answer in {gen_time:.2f}s")
                    logger.info(f"Sources: {sources}")
                    logger.info(f"Similarity scores: {similarity_scores_str}")
                    logger.info(f"Original Flesch: {original_flesch:.1f}, Middle School: {middle_school_flesch:.1f}, High School: {high_school_flesch:.1f}")
                
                qa_pairs.append((question, answer, sources, question_group, original_flesch, 
                               middle_school_answer, middle_school_flesch, 
                               high_school_answer, high_school_flesch, improved_answer, similarity_scores_str))
                
                # Write incrementally to CSV after each question
                self.write_csv([(question, answer, sources, question_group, original_flesch, 
                               middle_school_answer, middle_school_flesch, 
                               high_school_answer, high_school_flesch, improved_answer, similarity_scores_str)], 
                             kwargs.get('output_file', 'results.csv'), append=True, improved_prompt_text=improved_prompt_text)
                logger.info(f"Progress saved: {i+1}/{len(questions)} questions completed")
                
            except Exception as e:
                logger.error(f"Error processing question {i+1}: {e}")
                error_answer = "I encountered an error processing this question."
                sources = "Error retrieving sources"
                question_group = self._categorize_question(question)
                original_flesch = 0.0
                middle_school_answer = "I encountered an error processing this question."
                high_school_answer = "I encountered an error processing this question."
                improved_answer = "I encountered an error processing this question."
                middle_school_flesch = 0.0
                high_school_flesch = 0.0
                similarity_scores_str = "0.000"
                qa_pairs.append((question, error_answer, sources, question_group, original_flesch,
                               middle_school_answer, middle_school_flesch,
                               high_school_answer, high_school_flesch, improved_answer, similarity_scores_str))
                
                # Still write the error to CSV
                self.write_csv([(question, error_answer, sources, question_group, original_flesch,
                               middle_school_answer, middle_school_flesch,
                               high_school_answer, high_school_flesch, improved_answer, similarity_scores_str)], 
                             kwargs.get('output_file', 'results.csv'), append=True, improved_prompt_text=improved_prompt_text)
                logger.info(f"Error saved: {i+1}/{len(questions)} questions completed")
        
        return qa_pairs
    
    def _clean_readability_answer(self, answer: str, target_level: str) -> str:
        """Clean up readability-enhanced answers to remove unwanted phrases and formatting
        
        Args:
            answer: The readability-enhanced answer
            target_level: Either "middle_school" or "high_school"
        """
        cleaned = answer
        
        # Remove the "Here's a rewritten version" phrases
        if target_level == "middle_school":
            unwanted_phrases = [
                "Here's a rewritten version of the text at a middle school reading level:",
                "Here's a rewritten version of the text at a middle school reading level",
                "Here is a rewritten version of the text at a middle school reading level:",
                "Here is a rewritten version of the text at a middle school reading level",
                "Here's a rewritten version at a middle school reading level:",
                "Here's a rewritten version at a middle school reading level",
            ]
        elif target_level == "high_school":
            unwanted_phrases = [
                "Here's a rewritten version of the text at a high school reading level",
                "Here's a rewritten version of the text at a high school reading level:",
                "Here is a rewritten version of the text at a high school reading level",
                "Here is a rewritten version of the text at a high school reading level:",
                "Here's a rewritten version at a high school reading level",
                "Here's a rewritten version at a high school reading level:",
            ]
        else:
            unwanted_phrases = []
        
        for phrase in unwanted_phrases:
            if phrase.lower() in cleaned.lower():
                # Find and remove the phrase (case-insensitive)
                pattern = re.compile(re.escape(phrase), re.IGNORECASE)
                cleaned = pattern.sub("", cleaned).strip()
                # Remove leading colons, semicolons, or dashes
                cleaned = re.sub(r'^[:;\-]\s*', '', cleaned).strip()
        
        # Remove asterisks (but preserve bullet points if they use •)
        cleaned = re.sub(r'\*\*', '', cleaned)  # Remove bold markers
        cleaned = re.sub(r'\(\*\)', '', cleaned)  # Remove (*)
        cleaned = re.sub(r'\*', '', cleaned)  # Remove remaining asterisks
        
        # Clean up extra whitespace
        cleaned = ' '.join(cleaned.split())
        
        return cleaned
    
    def _clean_improved_answer(self, answer: str) -> str:
        """Clean up improved answer to remove unwanted phrases and formatting"""
        # Remove phrases like "Here's a rewritten version" or similar
        unwanted_phrases = [
            "Here's a rewritten version",
            "Here's a version",
            "Here is a rewritten version",
            "Here is a version",
            "Here's the answer",
            "Here is the answer"
        ]
        
        cleaned = answer
        for phrase in unwanted_phrases:
            if phrase.lower() in cleaned.lower():
                # Find and remove the phrase and any following colon/semicolon
                pattern = re.compile(re.escape(phrase), re.IGNORECASE)
                cleaned = pattern.sub("", cleaned).strip()
                # Remove leading colons, semicolons, or dashes
                cleaned = re.sub(r'^[:;\-]\s*', '', cleaned).strip()
        
        # Remove formatting markers like (*) or ** but preserve bullet points
        cleaned = re.sub(r'\*\*', '', cleaned)  # Remove bold markers
        cleaned = re.sub(r'\(\*\)', '', cleaned)  # Remove (*)
        # Note: Single asterisks are left alone as they might be used for formatting
        # The prompt specifies using • for bullet points, so this should be fine
        
        # Remove "Don't worry" and similar emotional management phrases
        emotional_phrases = [
            r"don't worry[^.]*\.\s*",
            r"Don't worry[^.]*\.\s*",
            r"the good news is[^.]*\.\s*",
            r"The good news is[^.]*\.\s*",
        ]
        for pattern in emotional_phrases:
            cleaned = re.sub(pattern, '', cleaned, flags=re.IGNORECASE)
        
        # Clean up extra whitespace
        cleaned = ' '.join(cleaned.split())
        
        return cleaned
    
    def diagnose_system(self, sample_questions: List[str] = None) -> Dict[str, Any]:
        """Diagnose the document loading, chunking, and retrieval system
        
        Args:
            sample_questions: Optional list of questions to test retrieval
            
        Returns:
            Dictionary with diagnostic information
        """
        diagnostics = {
            'vector_db_stats': {},
            'document_stats': {},
            'chunk_stats': {},
            'retrieval_tests': []
        }
        
        # Check vector database
        try:
            stats = self.vector_retriever.get_collection_stats()
            diagnostics['vector_db_stats'] = {
                'total_chunks': stats.get('total_chunks', 0),
                'collection_name': stats.get('collection_name', 'unknown'),
                'status': 'OK' if stats.get('total_chunks', 0) > 0 else 'EMPTY'
            }
        except Exception as e:
            diagnostics['vector_db_stats'] = {
                'status': 'ERROR',
                'error': str(e)
            }
        
        # Test document loading (without actually loading)
        try:
            data_path = Path(self.args.data_dir)
            if data_path.exists():
                supported_extensions = {'.txt', '.md', '.json', '.csv'}
                if PDF_AVAILABLE:
                    supported_extensions.add('.pdf')
                if DOCX_AVAILABLE:
                    supported_extensions.add('.docx')
                    supported_extensions.add('.doc')
                
                files = []
                for ext in supported_extensions:
                    files.extend(data_path.rglob(f"*{ext}"))
                
                # Sample a few files to check content
                sample_files = files[:5] if len(files) > 5 else files
                file_samples = []
                for file_path in sample_files:
                    try:
                        content = self._read_file(file_path)
                        file_samples.append({
                            'filename': file_path.name,
                            'size_chars': len(content),
                            'size_words': len(content.split()),
                            'readable': True
                        })
                    except Exception as e:
                        file_samples.append({
                            'filename': file_path.name,
                            'readable': False,
                            'error': str(e)
                        })
                
                diagnostics['document_stats'] = {
                    'total_files_found': len(files),
                    'sample_files': file_samples,
                    'status': 'OK'
                }
            else:
                diagnostics['document_stats'] = {
                    'status': 'ERROR',
                    'error': f'Data directory {self.args.data_dir} does not exist'
                }
        except Exception as e:
            diagnostics['document_stats'] = {
                'status': 'ERROR',
                'error': str(e)
            }
        
        # Test chunking on a sample document
        try:
            if diagnostics['document_stats'].get('status') == 'OK':
                sample_file = None
                for file_info in diagnostics['document_stats'].get('sample_files', []):
                    if file_info.get('readable', False):
                        # Find the actual file
                        data_path = Path(self.args.data_dir)
                        for ext in ['.txt', '.md', '.pdf', '.docx']:
                            files = list(data_path.rglob(f"*{file_info['filename']}"))
                            if files:
                                sample_file = files[0]
                                break
                        if sample_file:
                            break
                
                if sample_file:
                    content = self._read_file(sample_file)
                    # Create a dummy document (Document is already imported at top)
                    sample_doc = Document(
                        filename=sample_file.name,
                        content=content,
                        filepath=str(sample_file),
                        file_type=sample_file.suffix.lower(),
                        file_hash=""
                    )
                    
                    # Test chunking
                    sample_chunks = self._chunk_text(
                        content, 
                        sample_file.name, 
                        self.args.chunk_size, 
                        self.args.chunk_overlap
                    )
                    
                    chunk_lengths = [len(chunk.text.split()) for chunk in sample_chunks]
                    
                    diagnostics['chunk_stats'] = {
                        'sample_document': sample_file.name,
                        'total_chunks': len(sample_chunks),
                        'avg_chunk_size_words': sum(chunk_lengths) / len(chunk_lengths) if chunk_lengths else 0,
                        'min_chunk_size_words': min(chunk_lengths) if chunk_lengths else 0,
                        'max_chunk_size_words': max(chunk_lengths) if chunk_lengths else 0,
                        'chunk_size_setting': self.args.chunk_size,
                        'chunk_overlap_setting': self.args.chunk_overlap,
                        'status': 'OK'
                    }
        except Exception as e:
            diagnostics['chunk_stats'] = {
                'status': 'ERROR',
                'error': str(e)
            }
        
        # Test retrieval with sample questions
        if sample_questions and diagnostics['vector_db_stats'].get('status') == 'OK':
            for question in sample_questions:
                try:
                    context_chunks = self.retrieve(question, self.args.k)
                    sources = self._extract_sources(context_chunks)
                    
                    # Get similarity scores
                    results = self.vector_retriever.search(question, self.args.k)
                    
                    # Get sample chunk text (first 200 chars of first chunk)
                    sample_chunk_text = context_chunks[0].text[:200] + "..." if context_chunks else "N/A"
                    
                    diagnostics['retrieval_tests'].append({
                        'question': question,
                        'chunks_retrieved': len(context_chunks),
                        'sources': sources,
                        'similarity_scores': [f"{score:.3f}" for _, score in results],
                        'sample_chunk_preview': sample_chunk_text,
                        'status': 'OK' if context_chunks else 'NO_RESULTS'
                    })
                except Exception as e:
                    diagnostics['retrieval_tests'].append({
                        'question': question,
                        'status': 'ERROR',
                        'error': str(e)
                    })
        
        return diagnostics
    
    def print_diagnostics(self, diagnostics: Dict[str, Any]) -> None:
        """Print diagnostic information in a readable format"""
        print("\n" + "="*80)
        print("SYSTEM DIAGNOSTICS")
        print("="*80)
        
        # Vector DB Stats
        print("\n📊 VECTOR DATABASE:")
        vdb = diagnostics.get('vector_db_stats', {})
        print(f"  Status: {vdb.get('status', 'UNKNOWN')}")
        print(f"  Total chunks: {vdb.get('total_chunks', 0)}")
        print(f"  Collection: {vdb.get('collection_name', 'unknown')}")
        if 'error' in vdb:
            print(f"  Error: {vdb['error']}")
        
        # Document Stats
        print("\n📄 DOCUMENT LOADING:")
        doc_stats = diagnostics.get('document_stats', {})
        print(f"  Status: {doc_stats.get('status', 'UNKNOWN')}")
        print(f"  Total files found: {doc_stats.get('total_files_found', 0)}")
        if 'sample_files' in doc_stats:
            print(f"  Sample files:")
            for file_info in doc_stats['sample_files']:
                if file_info.get('readable', False):
                    print(f"    ✓ {file_info['filename']}: {file_info.get('size_chars', 0):,} chars, {file_info.get('size_words', 0):,} words")
                else:
                    print(f"    ✗ {file_info['filename']}: {file_info.get('error', 'unreadable')}")
        if 'error' in doc_stats:
            print(f"  Error: {doc_stats['error']}")
        
        # Chunk Stats
        print("\n✂️  CHUNKING:")
        chunk_stats = diagnostics.get('chunk_stats', {})
        print(f"  Status: {chunk_stats.get('status', 'UNKNOWN')}")
        if chunk_stats.get('status') == 'OK':
            print(f"  Sample document: {chunk_stats.get('sample_document', 'N/A')}")
            print(f"  Total chunks from sample: {chunk_stats.get('total_chunks', 0)}")
            print(f"  Average chunk size: {chunk_stats.get('avg_chunk_size_words', 0):.1f} words")
            print(f"  Chunk size range: {chunk_stats.get('min_chunk_size_words', 0)} - {chunk_stats.get('max_chunk_size_words', 0)} words")
            print(f"  Settings: size={chunk_stats.get('chunk_size_setting', 0)}, overlap={chunk_stats.get('chunk_overlap_setting', 0)}")
        if 'error' in chunk_stats:
            print(f"  Error: {chunk_stats['error']}")
        
        # Retrieval Tests
        if diagnostics.get('retrieval_tests'):
            print("\n🔍 RETRIEVAL TESTS:")
            for test in diagnostics['retrieval_tests']:
                print(f"\n  Question: {test.get('question', 'N/A')}")
                print(f"  Status: {test.get('status', 'UNKNOWN')}")
                if test.get('status') == 'OK':
                    print(f"  Chunks retrieved: {test.get('chunks_retrieved', 0)}")
                    print(f"  Sources: {test.get('sources', 'N/A')}")
                    scores = test.get('similarity_scores', [])
                    if scores:
                        print(f"  Similarity scores: {', '.join(scores)}")
                        # Warn if scores are low
                        try:
                            score_values = [float(s) for s in scores]
                            if max(score_values) < 0.3:
                                print(f"  ⚠️  WARNING: Low similarity scores - retrieved chunks may not be very relevant")
                            elif max(score_values) < 0.5:
                                print(f"  ⚠️  NOTE: Moderate similarity - consider increasing --k or checking chunk quality")
                        except:
                            pass
                    if 'sample_chunk_preview' in test:
                        print(f"  Sample chunk preview: {test['sample_chunk_preview']}")
                elif 'error' in test:
                    print(f"  Error: {test['error']}")
        
        print("\n" + "="*80 + "\n")
    
    def _extract_sources(self, context_chunks: List[Chunk]) -> str:
        """Extract source document names from context chunks"""
        sources = []
        for chunk in context_chunks:
            # Debug: Print chunk filename if verbose
            if self.args.verbose:
                logger.info(f"Chunk filename: {chunk.filename}")
            
            # Extract filename from chunk attribute (not metadata)
            source = chunk.filename if hasattr(chunk, 'filename') and chunk.filename else 'Unknown source'
            # Clean up the source name
            if source.endswith('.pdf'):
                source = source[:-4]  # Remove .pdf extension
            elif source.endswith('.txt'):
                source = source[:-4]  # Remove .txt extension
            elif source.endswith('.md'):
                source = source[:-3]  # Remove .md extension
            
            sources.append(source)
        
        # Remove duplicates while preserving order
        unique_sources = []
        for source in sources:
            if source not in unique_sources:
                unique_sources.append(source)
        
        return "; ".join(unique_sources)
    
    def _categorize_question(self, question: str) -> str:
        """Categorize a question into one of 5 categories"""
        question_lower = question.lower()
        
        # Gene-Specific Recommendations
        if any(gene in question_lower for gene in ['msh2', 'mlh1', 'msh6', 'pms2', 'epcam', 'brca1', 'brca2']):
            if any(kw in question_lower for kw in ['screening', 'surveillance', 'prevention', 'recommendation', 'risk', 'cancer risk', 'steps', 'management']):
                return "Gene-Specific Recommendations"
        
        # Inheritance Patterns
        if any(kw in question_lower for kw in ['inherit', 'inherited', 'pass', 'skip a generation', 'generation', 'can i pass']):
            return "Inheritance Patterns"
        
        # Family Risk Assessment
        if any(kw in question_lower for kw in ['family member', 'relative', 'first-degree', 'family risk', 'which relative', 'should my family']):
            return "Family Risk Assessment"
        
        # Genetic Variant Interpretation
        if any(kw in question_lower for kw in ['what does', 'genetic variant mean', 'variant mean', 'mutation mean', 'genetic result']):
            return "Genetic Variant Interpretation"
        
        # Support and Resources
        if any(kw in question_lower for kw in ['cope', 'overwhelmed', 'resource', 'genetic counselor', 'support', 'research', 'help', 'insurance', 'gina']):
            return "Support and Resources"
        
        # Default to Genetic Variant Interpretation if unclear
        return "Genetic Variant Interpretation"
    
    def enhance_readability(self, answer: str, target_level: str = "middle_school") -> Tuple[str, float]:
        """Enhance answer readability to different levels and calculate Flesch-Kincaid Grade Level
        
        Args:
            answer: The original answer to simplify
            target_level: One of "middle_school" or "high_school"
        
        Returns:
            Tuple of (simplified_answer, grade_level)
        """
        try:
            # Define prompts for different reading levels
            if target_level == "middle_school":
                level_description = "middle school reading level (ages 12-14, 6th-8th grade)"
                instructions = """
- Use simpler medical terms or explain them
- Medium-length sentences
- Clear, structured explanations
- Keep important medical information accessible"""
            elif target_level == "high_school":
                level_description = "high school reading level (ages 15-18, 9th-12th grade)"
                instructions = """
- Use appropriate medical terminology with context
- Varied sentence length
- Comprehensive yet accessible explanations
- Maintain technical accuracy while ensuring clarity"""
            else:
                raise ValueError(f"Unknown target_level: {target_level}")
            
            # Create a prompt to simplify the medical answer
            readability_prompt = f"""<|begin_of_text|><|start_header_id|>system<|end_header_id|>

You are a helpful medical assistant who specializes in explaining complex medical information at appropriate reading levels. Rewrite the following medical answer for {level_description}:
{instructions}
- Keep the same important information but adapt the complexity
- Provide context for technical terms
- Ensure the answer is informative yet understandable

<|eot_id|><|start_header_id|>user<|end_header_id|>

Please rewrite this medical answer for {level_description}:

{answer}<|eot_id|><|start_header_id|>assistant<|end_header_id|>

"""
            
            # Generate simplified answer
            inputs = self.tokenizer(readability_prompt, return_tensors="pt", truncation=True, max_length=2048)
            if self.device == "mps":
                inputs = {k: v.to(self.device) for k, v in inputs.items()}
            
            with torch.no_grad():
                outputs = self.model.generate(
                    **inputs,
                    max_new_tokens=512,  # Shorter for simplified version
                    temperature=0.3,     # Lower temperature for more consistent simplification
                    top_p=0.9,
                    repetition_penalty=1.05,
                    do_sample=True,
                    pad_token_id=self.tokenizer.eos_token_id,
                    eos_token_id=self.tokenizer.eos_token_id,
                    use_cache=True,
                    num_beams=1
                )
            
            # Decode response
            response = self.tokenizer.decode(outputs[0], skip_special_tokens=False)
            
            # Extract simplified answer
            prompt_end_marker = "<|start_header_id|>assistant<|end_header_id|>\n\n"
            if prompt_end_marker in response:
                simplified_answer = response.split(prompt_end_marker)[-1].strip()
            else:
                simplified_answer = response[len(readability_prompt):].strip()
            
            # Clean up special tokens
            if "<|eot_id|>" in simplified_answer:
                if simplified_answer.endswith("<|eot_id|>"):
                    simplified_answer = simplified_answer[:-len("<|eot_id|>")].strip()
            if "<|end_of_text|>" in simplified_answer:
                if simplified_answer.endswith("<|end_of_text|>"):
                    simplified_answer = simplified_answer[:-len("<|end_of_text|>")].strip()
            
            # Clean up unwanted phrases and formatting
            simplified_answer = self._clean_readability_answer(simplified_answer, target_level)
            
            # Calculate Flesch-Kincaid Grade Level
            try:
                grade_level = textstat.flesch_kincaid_grade(simplified_answer)
            except:
                grade_level = 0.0
            
            if self.args.verbose:
                logger.info(f"Simplified answer length: {len(simplified_answer)} characters")
                logger.info(f"Flesch-Kincaid Grade Level: {grade_level:.1f}")
            
            return simplified_answer, grade_level
            
        except Exception as e:
            logger.error(f"Error enhancing readability: {e}")
            # Fallback: return original answer with estimated grade level
            try:
                grade_level = textstat.flesch_kincaid_grade(answer)
            except:
                grade_level = 12.0  # Default to high school level
            return answer, grade_level
    
    def write_csv(self, qa_pairs: List[Tuple[str, str, str, str, float, str, float, str, float, str, str]], output_path: str, append: bool = False, improved_prompt_text: str = "") -> None:
        """Write Q&A pairs to CSV file in results folder
        
        Expected tuple format: (question, answer, sources, question_group, original_flesch, 
                               middle_school_answer, middle_school_flesch, 
                               high_school_answer, high_school_flesch, improved_answer, similarity_scores)
        """
        # Ensure results directory exists
        os.makedirs('results', exist_ok=True)
        
        # If output_path doesn't already have results/ prefix, add it
        if not output_path.startswith('results/'):
            output_path = f'results/{output_path}'
        
        if append:
            logger.info(f"Appending results to {output_path}")
        else:
            logger.info(f"Writing results to {output_path}")
        
        # Create output directory if needed
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        try:
            # Check if file exists and if we're appending
            file_exists = output_path.exists()
            write_mode = 'a' if append and file_exists else 'w'
            
            with open(output_path, write_mode, newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                
                # Write header only if creating new file or first append
                if not append or not file_exists:
                    # Create improved answer header with prompt text
                    improved_header = f'improved_answer (PROMPT: {improved_prompt_text})'
                    writer.writerow(['question', 'question_group', 'answer', 'original_flesch', 'sources', 
                                   'similarity_scores', 'middle_school_answer', 'middle_school_flesch', 
                                   'high_school_answer', 'high_school_flesch', improved_header])
                
                for data in qa_pairs:
                    # Unpack the data tuple
                    (question, answer, sources, question_group, original_flesch, 
                     middle_school_answer, middle_school_flesch, 
                     high_school_answer, high_school_flesch, improved_answer, similarity_scores) = data
                    
                    # Clean and escape the answers for CSV
                    def clean_text(text):
                        # Replace newlines with spaces and clean up formatting
                        cleaned = text.replace('\n', ' ').replace('\r', ' ')
                        # Remove extra whitespace but preserve the full content
                        cleaned = ' '.join(cleaned.split())
                        # Escape quotes properly for CSV
                        cleaned = cleaned.replace('"', '""')
                        return cleaned
                    
                    clean_question = clean_text(question)
                    clean_answer = clean_text(answer)
                    clean_sources = clean_text(sources)
                    clean_middle_school = clean_text(middle_school_answer)
                    clean_high_school = clean_text(high_school_answer)
                    clean_improved = clean_text(improved_answer)
                    
                    # Log the full answer length for debugging
                    if self.args.verbose:
                        logger.info(f"Writing answer length: {len(clean_answer)} characters")
                        logger.info(f"Middle school answer length: {len(clean_middle_school)} characters")
                        logger.info(f"High school answer length: {len(clean_high_school)} characters")
                        logger.info(f"Improved answer length: {len(clean_improved)} characters")
                        logger.info(f"Question group: {question_group}")
                    
                    # Use proper CSV quoting - let csv.writer handle the quoting
                    writer.writerow([
                        clean_question, 
                        question_group,
                        clean_answer, 
                        f"{original_flesch:.1f}",
                        clean_sources,
                        similarity_scores,  # Similarity scores (comma-separated)
                        clean_middle_school, 
                        f"{middle_school_flesch:.1f}",
                        clean_high_school, 
                        f"{high_school_flesch:.1f}",
                        clean_improved
                    ])
            
            if append:
                logger.info(f"Appended {len(qa_pairs)} Q&A pairs to {output_path}")
            else:
                logger.info(f"Successfully wrote {len(qa_pairs)} Q&A pairs to {output_path}")
            
        except Exception as e:
            logger.error(f"Failed to write CSV: {e}")
            sys.exit(4)


def parse_args():
    """Parse command line arguments"""
    parser = argparse.ArgumentParser(description="RAG Chatbot for CGT-LLM-Beta with Vector Database")
    
    # File paths
    parser.add_argument('--data-dir', default='./Data Resources', 
                       help='Directory containing documents to index')
    parser.add_argument('--questions', default='./questions.txt',
                       help='File containing questions (one per line)')
    parser.add_argument('--out', default='./answers.csv',
                       help='Output CSV file for answers')
    parser.add_argument('--vector-db-dir', default='./chroma_db',
                       help='Directory for ChromaDB persistence')
    
    # Retrieval parameters
    parser.add_argument('--k', type=int, default=5,
                       help='Number of chunks to retrieve per question')
    
    # Chunking parameters
    parser.add_argument('--chunk-size', type=int, default=500,
                       help='Size of text chunks in tokens')
    parser.add_argument('--chunk-overlap', type=int, default=200,
                       help='Overlap between chunks in tokens')
    
    # Model selection
    parser.add_argument('--model', type=str, default='meta-llama/Llama-3.2-3B-Instruct',
                       help='HuggingFace model name to use (e.g., meta-llama/Llama-3.2-3B-Instruct, mistralai/Mistral-7B-Instruct-v0.2)')
    
    # Generation parameters
    parser.add_argument('--max-new-tokens', type=int, default=1024,
                       help='Maximum new tokens to generate')
    parser.add_argument('--temperature', type=float, default=0.2,
                       help='Generation temperature')
    parser.add_argument('--top-p', type=float, default=0.9,
                       help='Top-p sampling parameter')
    parser.add_argument('--repetition-penalty', type=float, default=1.1,
                       help='Repetition penalty')
    
    # Database options
    parser.add_argument('--force-rebuild', action='store_true',
                       help='Force rebuild of vector database')
    parser.add_argument('--skip-indexing', action='store_true',
                       help='Skip document indexing, use existing database')
    
    # Other options
    parser.add_argument('--seed', type=int, default=42,
                       help='Random seed for reproducibility')
    parser.add_argument('--verbose', action='store_true',
                       help='Enable verbose logging')
    parser.add_argument('--dry-run', action='store_true',
                       help='Build index and test retrieval without generation')
    parser.add_argument('--diagnose', action='store_true',
                       help='Run system diagnostics and exit')
    
    return parser.parse_args()


def main():
    """Main function"""
    args = parse_args()
    
    # Set random seed
    torch.manual_seed(args.seed)
    np.random.seed(args.seed)
    
    # Set logging level
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    logger.info("Starting RAG Chatbot with Vector Database")
    logger.info(f"Arguments: {vars(args)}")
    
    try:
        # Initialize bot
        bot = RAGBot(args)
        
        # Check if we should skip indexing
        if not args.skip_indexing:
            # Load and process documents
            documents = bot.load_corpus(args.data_dir)
            if not documents:
                logger.error("No documents found to process")
                sys.exit(3)
            
            # Chunk documents
            chunks = bot.chunk_documents(documents, args.chunk_size, args.chunk_overlap)
            if not chunks:
                logger.error("No chunks created from documents")
                sys.exit(3)
            
            # Build or update index
            bot.build_or_update_index(chunks, args.force_rebuild)
        else:
            logger.info("Skipping document indexing, using existing vector database")
        
        # Run diagnostics if requested
        if args.diagnose:
            sample_questions = [
                "What is Lynch Syndrome?",
                "What does a BRCA1 genetic variant mean?",
                "What screening tests are recommended for MSH2 carriers?"
            ]
            diagnostics = bot.diagnose_system(sample_questions=sample_questions)
            bot.print_diagnostics(diagnostics)
            return
        
        if args.dry_run:
            logger.info("Dry run completed successfully")
            return
        
        # Process questions
        generation_kwargs = {
            'max_new_tokens': args.max_new_tokens,
            'temperature': args.temperature,
            'top_p': args.top_p,
            'repetition_penalty': args.repetition_penalty
        }
        
        qa_pairs = bot.process_questions(args.questions, output_file=args.out, **generation_kwargs)
        
        logger.info("RAG Chatbot completed successfully")
        
    except KeyboardInterrupt:
        logger.info("Interrupted by user")
        sys.exit(0)
    except Exception as e:
        logger.error(f"Unexpected error: {e}")
        if args.verbose:
            import traceback
            traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()